"""
水印服务 - 支持图片、PDF、Word文档
"""
import io
import glob
import math
import logging
from pathlib import Path
from typing import Optional, Tuple
from PIL import Image, ImageDraw, ImageFont, ImageEnhance
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config as app_config
from models import WatermarkConfig, WatermarkPosition


def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """HEX颜色转RGB"""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def hex_to_rgba(hex_color: str, opacity: float) -> Tuple[int, int, int, int]:
    """HEX颜色转RGBA"""
    r, g, b = hex_to_rgb(hex_color)
    a = int(opacity * 255)
    return (r, g, b, a)


# 全局字体路径配置（支持中文）
CHINESE_FONT_PATHS = [
    # Docker 容器内置字体（优先，TTF格式兼容reportlab）
    "/usr/share/fonts/chinese/LXGWWenKai-Regular.ttf",  # 霞鹜文楷
    # Windows
    "C:/Windows/Fonts/msyh.ttc",      # 微软雅黑
    "C:/Windows/Fonts/msyhbd.ttc",    # 微软雅黑粗体
    "C:/Windows/Fonts/simhei.ttf",    # 黑体
    "C:/Windows/Fonts/simsun.ttc",    # 宋体
    "C:/Windows/Fonts/simkai.ttf",    # 楷体
    "C:/Windows/Fonts/STZHONGS.TTF",  # 华文中宋
    # macOS
    "/System/Library/Fonts/PingFang.ttc",           # 苹方
    "/System/Library/Fonts/STHeiti Light.ttc",      # 华文黑体
    "/Library/Fonts/Arial Unicode.ttf",
]


# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _find_cjk_fonts() -> list:
    """动态查找系统中的 CJK 字体"""
    patterns = [
        "/usr/share/fonts/**/*.otf",
        "/usr/share/fonts/**/*.ttf",
        "/usr/share/fonts/**/*.ttc",
    ]
    found = []
    for pattern in patterns:
        found.extend(glob.glob(pattern, recursive=True))
    logger.info(f"动态查找到的字体: {found}")
    return found


def get_font(size: int) -> ImageFont.FreeTypeFont:
    """获取字体，优先使用系统中文字体"""
    # 优先使用自定义字体路径
    font_paths = []
    if app_config.CUSTOM_FONT_PATH and Path(app_config.CUSTOM_FONT_PATH).exists():
        font_paths.append(app_config.CUSTOM_FONT_PATH)
    font_paths.extend(CHINESE_FONT_PATHS)
    # 动态查找的字体作为后备
    font_paths.extend(_find_cjk_fonts())
    
    logger.info(f"尝试加载字体，候选列表: {font_paths}")
    
    for font_path in font_paths:
        exists = Path(font_path).exists()
        logger.info(f"检查字体: {font_path}, 存在: {exists}")
        if exists:
            try:
                # .ttc文件需要指定字体索引
                if font_path.lower().endswith('.ttc'):
                    font = ImageFont.truetype(font_path, size, index=0)
                else:
                    font = ImageFont.truetype(font_path, size)
                logger.info(f"成功加载字体: {font_path}")
                return font
            except Exception as e:
                logger.error(f"加载字体失败 {font_path}: {e}")
                continue
    
    # 降级使用默认字体
    logger.warning("未找到中文字体，使用默认字体")
    return ImageFont.load_default()


class ImageWatermarker:
    """图片水印处理器"""
    
    @staticmethod
    def add_watermark(
        image_data: bytes,
        text: str,
        config: WatermarkConfig,
        output_format: str = "PNG"
    ) -> bytes:
        """为图片添加水印"""
        # 打开图片
        img = Image.open(io.BytesIO(image_data))
        
        # 转换为RGBA模式以支持透明度
        if img.mode != "RGBA":
            img = img.convert("RGBA")
        
        # 创建水印层
        watermark_layer = Image.new("RGBA", img.size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(watermark_layer)
        
        # 获取字体
        font = get_font(config.font_size)
        
        # 获取颜色
        color = hex_to_rgba(config.font_color, config.opacity)
        
        # 获取文字大小
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        if config.position == WatermarkPosition.TILE:
            # 平铺水印
            ImageWatermarker._add_tile_watermark(
                draw, img.size, text, font, color,
                text_width, text_height, config.spacing, config.angle
            )
        elif config.position == WatermarkPosition.CENTER:
            # 居中水印
            x = (img.size[0] - text_width) // 2
            y = (img.size[1] - text_height) // 2
            draw.text((x, y), text, font=font, fill=color)
        else:
            # 角落水印
            x, y = ImageWatermarker._get_corner_position(
                img.size, text_width, text_height, config.position
            )
            draw.text((x, y), text, font=font, fill=color)
        
        # 旋转水印层（如果不是平铺模式）
        if config.position != WatermarkPosition.TILE and config.angle != 0:
            watermark_layer = watermark_layer.rotate(
                config.angle, resample=Image.BICUBIC, expand=False
            )
        
        # 合并图层
        result = Image.alpha_composite(img, watermark_layer)
        
        # 如果原图是RGB模式，转换回RGB
        if output_format.upper() in ["JPEG", "JPG"]:
            result = result.convert("RGB")
        
        # 保存到字节流
        output = io.BytesIO()
        save_format = "JPEG" if output_format.upper() in ["JPEG", "JPG"] else output_format.upper()
        result.save(output, format=save_format, quality=95)
        return output.getvalue()
    
    @staticmethod
    def _add_tile_watermark(
        draw: ImageDraw.Draw,
        img_size: Tuple[int, int],
        text: str,
        font: ImageFont.FreeTypeFont,
        color: Tuple[int, int, int, int],
        text_width: int,
        text_height: int,
        spacing: int,
        angle: float
    ):
        """添加平铺水印"""
        # 计算旋转后需要覆盖的区域
        diagonal = math.sqrt(img_size[0]**2 + img_size[1]**2)
        
        # 创建临时图层用于旋转
        temp_size = int(diagonal * 1.5)
        temp_layer = Image.new("RGBA", (temp_size, temp_size), (255, 255, 255, 0))
        temp_draw = ImageDraw.Draw(temp_layer)
        
        step_x = text_width + spacing
        step_y = text_height + spacing
        
        for y in range(0, temp_size, step_y):
            for x in range(0, temp_size, step_x):
                temp_draw.text((x, y), text, font=font, fill=color)
        
        # 旋转
        temp_layer = temp_layer.rotate(angle, resample=Image.BICUBIC, expand=False)
        
        # 裁剪到原图大小
        center_x = temp_size // 2
        center_y = temp_size // 2
        left = center_x - img_size[0] // 2
        top = center_y - img_size[1] // 2
        right = left + img_size[0]
        bottom = top + img_size[1]
        
        cropped = temp_layer.crop((left, top, right, bottom))
        
        # 将裁剪后的水印绘制到原draw上
        # 由于我们需要返回给原始draw，这里改用另一种方式
        # 直接在原始watermark_layer上绘制
        draw._image.paste(cropped, (0, 0), cropped)
    
    @staticmethod
    def _get_corner_position(
        img_size: Tuple[int, int],
        text_width: int,
        text_height: int,
        position: WatermarkPosition
    ) -> Tuple[int, int]:
        """获取角落位置坐标"""
        margin = 20
        positions = {
            WatermarkPosition.TOP_LEFT: (margin, margin),
            WatermarkPosition.TOP_RIGHT: (img_size[0] - text_width - margin, margin),
            WatermarkPosition.BOTTOM_LEFT: (margin, img_size[1] - text_height - margin),
            WatermarkPosition.BOTTOM_RIGHT: (
                img_size[0] - text_width - margin,
                img_size[1] - text_height - margin
            ),
        }
        return positions.get(position, (margin, margin))


class PDFWatermarker:
    """PDF水印处理器"""
    
    @staticmethod
    def add_watermark(
        pdf_data: bytes,
        text: str,
        config: WatermarkConfig
    ) -> bytes:
        """为PDF添加水印"""
        # 创建水印PDF
        watermark_pdf = PDFWatermarker._create_watermark_pdf(text, config)
        
        # 读取原PDF
        reader = PdfReader(io.BytesIO(pdf_data))
        writer = PdfWriter()
        
        # 读取水印PDF
        watermark_reader = PdfReader(io.BytesIO(watermark_pdf))
        watermark_page = watermark_reader.pages[0]
        
        # 为每一页添加水印
        for page in reader.pages:
            page.merge_page(watermark_page)
            writer.add_page(page)
        
        # 输出
        output = io.BytesIO()
        writer.write(output)
        return output.getvalue()
    
    @staticmethod
    def _create_watermark_pdf(text: str, config: WatermarkConfig) -> bytes:
        """创建水印PDF"""
        output = io.BytesIO()
        c = canvas.Canvas(output, pagesize=letter)
        width, height = letter
        
        # 设置透明度
        c.setFillAlpha(config.opacity)
        
        # 设置颜色
        r, g, b = hex_to_rgb(config.font_color)
        c.setFillColorRGB(r/255, g/255, b/255)
        
        # 注册中文字体
        font_name = "Helvetica"
        font_registered = False
        
        # 优先使用自定义字体路径
        font_paths = []
        if app_config.CUSTOM_FONT_PATH and Path(app_config.CUSTOM_FONT_PATH).exists():
            font_paths.append(app_config.CUSTOM_FONT_PATH)
        font_paths.extend(CHINESE_FONT_PATHS)
        font_paths.extend(_find_cjk_fonts())
        
        for font_path in font_paths:
            if Path(font_path).exists() and not font_registered:
                try:
                    # 生成唯一的字体名称
                    path_obj = Path(font_path)
                    name = path_obj.stem.lower().replace(' ', '_')
                    
                    # .ttc文件需要特殊处理，使用subfontIndex
                    if font_path.lower().endswith('.ttc'):
                        pdfmetrics.registerFont(TTFont(name, font_path, subfontIndex=0))
                    else:
                        pdfmetrics.registerFont(TTFont(name, font_path))
                    
                    font_name = name
                    font_registered = True
                    break
                except Exception:
                    continue
        
        c.setFont(font_name, config.font_size)
        
        if config.position == WatermarkPosition.TILE:
            # 平铺水印
            spacing = config.spacing + config.font_size * len(text) * 0.6
            c.saveState()
            c.translate(width/2, height/2)
            c.rotate(config.angle)
            
            for y in range(-int(height), int(height), int(spacing)):
                for x in range(-int(width), int(width), int(spacing)):
                    c.drawString(x, y, text)
            
            c.restoreState()
        elif config.position == WatermarkPosition.CENTER:
            c.saveState()
            c.translate(width/2, height/2)
            c.rotate(config.angle)
            c.drawCentredString(0, 0, text)
            c.restoreState()
        else:
            # 角落位置
            x, y = PDFWatermarker._get_corner_position(
                width, height, config.font_size, len(text), config.position
            )
            c.drawString(x, y, text)
        
        c.save()
        return output.getvalue()
    
    @staticmethod
    def _get_corner_position(
        width: float,
        height: float,
        font_size: int,
        text_len: int,
        position: WatermarkPosition
    ) -> Tuple[float, float]:
        """获取角落位置"""
        margin = 50
        text_width = font_size * text_len * 0.6
        
        positions = {
            WatermarkPosition.TOP_LEFT: (margin, height - margin),
            WatermarkPosition.TOP_RIGHT: (width - text_width - margin, height - margin),
            WatermarkPosition.BOTTOM_LEFT: (margin, margin),
            WatermarkPosition.BOTTOM_RIGHT: (width - text_width - margin, margin),
        }
        return positions.get(position, (margin, margin))


class WordWatermarker:
    """Word文档水印处理器"""
    
    @staticmethod
    def add_watermark(
        docx_data: bytes,
        text: str,
        config: WatermarkConfig
    ) -> bytes:
        """为Word文档添加水印"""
        doc = Document(io.BytesIO(docx_data))
        
        # 为每个section添加水印
        for section in doc.sections:
            WordWatermarker._add_watermark_to_section(section, text, config)
        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    @staticmethod
    def _add_watermark_to_section(section, text: str, config: WatermarkConfig):
        """为section添加水印"""
        # 获取或创建header
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        
        paragraph = header.paragraphs[0]
        
        # 创建水印形状
        r, g, b = hex_to_rgb(config.font_color)
        
        # 使用VML创建文字水印
        pict = OxmlElement('w:pict')
        
        shape = OxmlElement('v:shape')
        shape.set(qn('id'), 'watermark')
        shape.set('type', '#_x0000_t136')  # 文字水印类型
        shape.set('style', f'position:absolute;margin-left:0;margin-top:0;width:500pt;height:200pt;rotation:{int(config.angle)};z-index:-251658752')
        shape.set('fillcolor', config.font_color)
        shape.set('stroked', 'f')
        
        # 设置填充透明度
        fill = OxmlElement('v:fill')
        fill.set('opacity', str(config.opacity))
        shape.append(fill)
        
        # 文字路径
        textpath = OxmlElement('v:textpath')
        textpath.set('style', f'font-family:"Microsoft YaHei";font-size:{config.font_size}pt')
        textpath.set('string', text)
        shape.append(textpath)
        
        pict.append(shape)
        
        # 将水印添加到段落
        run = paragraph.add_run()
        run._element.append(pict)


def add_watermark(
    file_data: bytes,
    text: str,
    file_type: str,
    config: Optional[WatermarkConfig] = None,
    file_extension: str = ".png"
) -> bytes:
    """
    统一的水印添加接口
    
    Args:
        file_data: 文件二进制数据
        text: 水印文字
        file_type: 文件类型 (image/pdf/word)
        config: 水印配置
        file_extension: 文件扩展名（用于图片格式判断）
    
    Returns:
        添加水印后的文件二进制数据
    """
    if config is None:
        config = WatermarkConfig()
    
    if file_type == "image":
        # 根据扩展名确定输出格式
        format_map = {
            ".jpg": "JPEG",
            ".jpeg": "JPEG",
            ".png": "PNG",
            ".gif": "GIF",
            ".bmp": "BMP",
            ".webp": "WEBP",
            ".tiff": "TIFF",
        }
        output_format = format_map.get(file_extension.lower(), "PNG")
        return ImageWatermarker.add_watermark(file_data, text, config, output_format)
    
    elif file_type == "pdf":
        return PDFWatermarker.add_watermark(file_data, text, config)
    
    elif file_type == "word":
        return WordWatermarker.add_watermark(file_data, text, config)
    
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
